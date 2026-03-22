"""Tests for document processing pipeline — upload → parse → chunk → embed → index."""

from __future__ import annotations

import pytest

from mcp_manager.documents.embeddings import LocalEmbeddingClient
from mcp_manager.documents.index import DocumentIndex
from mcp_manager.documents.pipeline import DocumentProcessor
from mcp_manager.documents.store import FileSystemDocumentStore


@pytest.fixture()
def pipeline(tmp_path):
    store = FileSystemDocumentStore(base_dir=str(tmp_path))
    index = DocumentIndex()
    embedder = LocalEmbeddingClient(dimension=64)
    return DocumentProcessor(
        store=store, index=index, embedding_client=embedder,
        chunk_size=50, chunk_overlap=10,
    ), index


class TestDocumentProcessor:
    @pytest.mark.asyncio
    async def test_process_text_file(self, pipeline):
        processor, index = pipeline
        content = b"This is a test document with enough words to verify chunking works correctly. " * 10

        doc = await processor.process("tenant-a", "docs", "test.txt", content)
        assert doc.status == "ready"
        assert doc.chunk_count > 0

        # Verify it's in the index
        docs = index.list_documents("tenant-a")
        assert len(docs) == 1

    @pytest.mark.asyncio
    async def test_process_markdown_file(self, pipeline):
        processor, index = pipeline
        content = b"# Title\n\nThis is markdown content with **bold** and _italic_ text. " * 5

        doc = await processor.process("tenant-a", "docs", "readme.md", content)
        assert doc.status == "ready"

    @pytest.mark.asyncio
    async def test_process_pdf_file(self, pipeline):
        processor, index = pipeline

        from io import BytesIO
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        buf = BytesIO()
        c = canvas.Canvas(buf, pagesize=letter)
        c.drawString(72, 700, "This is a PDF test document with important content about machine learning.")
        c.showPage()
        c.save()

        doc = await processor.process("tenant-a", "docs", "paper.pdf", buf.getvalue())
        assert doc.status == "ready"
        assert doc.chunk_count > 0

    @pytest.mark.asyncio
    async def test_process_docx_file(self, pipeline):
        processor, index = pipeline

        import docx
        from io import BytesIO

        d = docx.Document()
        d.add_paragraph("First paragraph about artificial intelligence.")
        d.add_paragraph("Second paragraph about neural networks and deep learning.")
        buf = BytesIO()
        d.save(buf)

        doc = await processor.process("tenant-a", "docs", "report.docx", buf.getvalue())
        assert doc.status == "ready"

    @pytest.mark.asyncio
    async def test_unsupported_format(self, pipeline):
        processor, index = pipeline
        doc = await processor.process("tenant-a", "docs", "data.xlsx", b"binary")
        assert doc.status == "error"
        assert "Unsupported" in (doc.error_message or "")

    @pytest.mark.asyncio
    async def test_tenant_isolation(self, pipeline):
        processor, index = pipeline

        await processor.process("tenant-a", "docs", "a.txt", b"Content for A " * 20)
        await processor.process("tenant-b", "docs", "b.txt", b"Content for B " * 20)

        assert len(index.list_documents("tenant-a")) == 1
        assert len(index.list_documents("tenant-b")) == 1

    @pytest.mark.asyncio
    async def test_searchable_after_processing(self, pipeline):
        processor, index = pipeline
        embedder = LocalEmbeddingClient(dimension=64)

        content = b"Machine learning algorithms are used for predictive modeling and data analysis. " * 10
        await processor.process("tenant-a", "docs", "ml.txt", content)

        query_emb = await embedder.embed_single("machine learning")
        results = index.search("tenant-a", query_emb)
        assert len(results) > 0
        assert "machine learning" in results[0].chunk_text.lower()
