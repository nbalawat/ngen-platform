"""Tests for document parser — creates real files, no mocks."""

from __future__ import annotations

import pytest

from mcp_manager.documents.parser import UnsupportedFormatError, parse_document


class TestTextParsing:
    def test_txt_file(self):
        text = parse_document("readme.txt", b"Hello, World!")
        assert text == "Hello, World!"

    def test_markdown_file(self):
        md = b"# Title\n\nSome **bold** text."
        text = parse_document("docs.md", md)
        assert "# Title" in text
        assert "**bold**" in text

    def test_utf8_with_bom(self):
        content = b"\xef\xbb\xbfHello UTF-8 BOM"
        text = parse_document("test.txt", content)
        assert "Hello UTF-8 BOM" in text


class TestPdfParsing:
    def test_pdf_extraction(self):
        """Create a real PDF and extract text from it."""
        from io import BytesIO
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        buf = BytesIO()
        c = canvas.Canvas(buf, pagesize=letter)
        c.drawString(72, 700, "This is a test PDF document.")
        c.drawString(72, 680, "It has multiple lines of text.")
        c.showPage()
        c.save()

        text = parse_document("test.pdf", buf.getvalue())
        assert "test PDF document" in text
        assert "multiple lines" in text

    def test_pdf_multipage(self):
        """PDF with multiple pages extracts all pages."""
        from io import BytesIO
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        buf = BytesIO()
        c = canvas.Canvas(buf, pagesize=letter)
        c.drawString(72, 700, "Page one content here.")
        c.showPage()
        c.drawString(72, 700, "Page two content here.")
        c.showPage()
        c.save()

        text = parse_document("multi.pdf", buf.getvalue())
        assert "Page one" in text
        assert "Page two" in text


class TestDocxParsing:
    def test_docx_extraction(self):
        """Create a real DOCX and extract text."""
        import docx
        from io import BytesIO

        doc = docx.Document()
        doc.add_paragraph("First paragraph of the document.")
        doc.add_paragraph("Second paragraph with important content.")
        buf = BytesIO()
        doc.save(buf)

        text = parse_document("test.docx", buf.getvalue())
        assert "First paragraph" in text
        assert "important content" in text


class TestUnsupportedFormat:
    def test_unknown_extension(self):
        with pytest.raises(UnsupportedFormatError, match="xlsx"):
            parse_document("data.xlsx", b"binary data")

    def test_no_extension(self):
        with pytest.raises(UnsupportedFormatError):
            parse_document("noext", b"some data")
